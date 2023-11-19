from flask import Flask, render_template, request, send_file
from flask_cors import CORS, cross_origin
import pandas as pd
from docx import Document

app = Flask(__name__)
CORS(app)


@cross_origin()
@app.route("/")
def index():
    print("index come!")
    return render_template('./index.html')


@app.route("/getFile2")
def xx():
    return "Asdfasdf2"


@app.route("/api/my", methods=['POST'])
def xx2():

    print("--------")
    print(request.method)
    # print)
    dfSrc = {
        '구분': [],
        '지원부서': [],
        '1차검토': [],
        "현업검토": [],
        "성명": [],
        '생년월일': [],
        '연락처': [],
        "메일주소": [],
        '주소': [],
        "최종학력": [],
        "전공": [],
        '졸업연월': [],
        "경력사항1": [],
        "최종직위": [],
        "경력사항2": [],
        '직위2': [],
        "경력사항3": [],
        '직위3': [],
    }
    if (request.method == "POST"):
        print('post come')
        files = request.files.getlist('files[]')
        for fileIdx, f in enumerate(files):
            for key in dfSrc.keys():
                dfSrc[key].append("")
            doc = Document(f)

            tableOne = doc.tables[0]
            for rowIdx, row in enumerate(tableOne.rows):
                text = (cell.text for cell in row.cells)
                listText = list(text)

                if (rowIdx == 0):
                    dfSrc["성명"][fileIdx] = listText[2]
                    dfSrc["지원부서"][fileIdx] = listText[6]
                if (rowIdx == 1):
                    dfSrc["생년월일"][fileIdx] = listText[2]
                if (rowIdx == 2):
                    dfSrc["연락처"][fileIdx] = listText[2]
                if (rowIdx == 3):
                    dfSrc["메일주소"][fileIdx] = listText[2]
                if (rowIdx == 4):
                    dfSrc["주소"][fileIdx] = listText[2]

            tableTwo = doc.tables[1]
            tableTwoTargetIdx = 0

            for rowIdx, row in enumerate(tableTwo.rows):
                text = (cell.text for cell in row.cells)
                listText = list(text)
                if (listText[1].strip() != ""):
                    tableTwoTargetIdx = rowIdx

            tableTwoTargetRow = list(cell.text for cell in (
                list(tableTwo.rows))[tableTwoTargetIdx].cells)
            dfSrc["졸업연월"][fileIdx] = tableTwoTargetRow[0]
            dfSrc["최종학력"][fileIdx] = tableTwoTargetRow[1]
            dfSrc["전공"][fileIdx] = tableTwoTargetRow[2]

            tableThree = doc.tables[2]

            tableTreeTargetIdx = 0

            for rowIdx, row in enumerate(tableThree.rows):

                text = (cell.text for cell in row.cells)
                listText = list(text)
                compName = listText[1]
                level = listText[3]
                if (rowIdx == 0):
                    continue
                if (rowIdx > 3):
                    continue
                if (len(compName) < 1):
                    continue
                if (rowIdx == 1):
                    dfSrc["경력사항1"][fileIdx] = compName
                    dfSrc["최종직위"][fileIdx] = level
                    continue
                if (rowIdx == 2):
                    dfSrc["경력사항2"][fileIdx] = compName
                    dfSrc["직위2"][fileIdx] = level
                    continue
                if (rowIdx == 3):
                    dfSrc["경력사항3"][fileIdx] = compName
                    dfSrc["직위3"][fileIdx] = level
                    continue

        df1 = pd.DataFrame(dfSrc)

        with pd.ExcelWriter('static/output.xlsx') as writer:
            df1.to_excel(writer, sheet_name='Sheet_name_1')
        try:
            return send_file('static/output.xlsx', as_attachment=True)
        except:
            print('fail..')
            return "fail.."


# def main2():
#     # app.run(host='0.0.0.0')
#     app.run(host='127.0.0.1', port=3070)


# if __name__ == "__main__":
#     main2()
